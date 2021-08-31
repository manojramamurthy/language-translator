def docx_to_lang(name_of_file, output_file=None, target=None, heading=None):
    from docx import Document
    from mtranslate import translate
    # Import docx NOT python-docx
    import docx
    # Create an instance of a word document
    doc = docx.Document()
    doc.add_heading(heading, 0)

    docx_dict = {}
    lang = []
    document = Document(name_of_file)
    idx = 0
    for para in document.paragraphs:
        idx+=1
        if(len(para.text)>0):
            docx_dict[idx] = para.text
            # lang.append(translate(para.text, target))
            doc.add_heading('--------------------**************************------------------------')
            doc.add_paragraph(translate(para.text, target), style='List Number')
    # Now save the document to a location 
    doc.save(output_file)
    return [docx_dict, lang]

# target = 'ta' is tamil
docx = docx_to_lang('sample/sample.docx', output_file='sample/sample-output-tamil.docx', target='ta', heading='Sample Heading')
print(docx[0])